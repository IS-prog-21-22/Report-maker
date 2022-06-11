import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from zipfile import ZipFile

from numpy import full

def main():
    print("Эта программа пробегается по папке, откуда запущена, ищет подпапки с именами типа 'lab_', где вместо",
    "нижнего подчёркивания стоит цифра от 1 до 6, и, используя неглубинные - т.е. родители которых это подпапки 'lab_' -",
    "файлы с расширениями .cpp, .hpp и .h, создаёт отчёты по лабам [для 1 курса ИСа ИТМО 2022]")
    print("NB! Программа поставляется as is, проверяйте отчёты после выполнения программы, только вы несёте за них ответственность!")
    full_name = input("Введите ФИО: ")
    group_number = input("Введите номер группы: ")
    gender = input("Введите гендер(ж/м): ")
    if gender not in ['м', 'ж']:
        print("Ошибка при вводе гендера")
        return

    lab_names_lst = ["ООП. Классы", "Использование внешних библиотек", "STL. Контейнеры", "Кубик Рубика", "Allocator", "Programming at compile-time"]
    cur_folder = os.getcwd()
    files_to_zip = []
    for i in range(1,7):
        sub_folder = os.path.join(cur_folder, f"lab{i}")
        if not os.path.exists(sub_folder):
            print(f"Не существует {sub_folder} - не будет соответствующего отчёта")
            continue

        document = Document()
        style = document.styles["Normal"]
        style.font.name = "Calibri"
        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run("Министерство науки и высшего образования Российской Федерации\n")
        run.bold = True
        run.font.size = Pt(18)
        run = p.add_run("Федеральное государственное автономное образовательное учреждение высшего образования " +
        '"Национальный исследовательский университет ИТМО"\n')
        run.bold = True
        run.font.size = Pt(18)
        run = p.add_run("Факультет информационных технологий и программирования\n\n")
        run.font.size = Pt(18)
        p.add_run(f"Лабораторная №{i}\n")
        p.add_run(f"{lab_names_lst[i - 1]}\n\n").italic = True

        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        if gender == 'ж':
            p.add_run(f"Выполнила студентка группы {group_number}\n").bold = True
        else:
            p.add_run(f"Выполнил студент группы {group_number}\n").bold = True
        p.add_run(full_name)

        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run("\n" * 18 + "Санкт-Петербург\n2022")

        for file in os.listdir(sub_folder):
            if any(map(lambda end: file.endswith(end), [".cpp", ".hpp", ".h"])):
                with open(os.path.join(sub_folder, file), 'r', encoding='utf-8') as input_file:
                    p = document.add_paragraph(f"Файл: {file}\n\n")
                    p.paragraph_format.page_break_before = True
                    for line in input_file.readlines():
                        p.add_run(line)

        document_name = "_".join(full_name.split()) + f"_{group_number}_Лабораторная_{i}.docx"
        document.save(document_name)
        files_to_zip.append(document_name)
    
    with ZipFile("_".join(full_name.split()) + f"_{group_number}.zip", "w") as output_file:
        for file in files_to_zip:
            output_file.write(file)
    
    for file in files_to_zip:
        os.remove(file)

    print("Успешно!")
    _ = input("Подтвердите завершение программы(Enter):")


if __name__ == "__main__":
    main()